import os

from pypdf import PdfReader
# from docx import Document
import pandas as pd
# from langchain_mistralai import ChatMistralAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document
# from langchain_huggingface import HuggingFaceEmbeddings
from docx import Document as DocxDocument
from langchain_core.documents import Document

# Read PDF
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    docs = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={"page": i, "source": file_path}
            ))
    return docs

# Read .docx
def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return [
        Document(page_content=para.text, metadata={"source": file_path})
        for para in doc.paragraphs if para.text
    ]


# Read Excel
def extract_text_from_excel(file_path):
    df = pd.read_excel(file_path)
    return [
        Document(page_content=df.to_string(), metadata={"source": file_path})
    ]

# Read CSV
def extract_text_from_csv(file_path):
    df = pd.read_csv(file_path)
    return [
        Document(page_content=df.to_string(), metadata={"source": file_path})
    ]

#Read .txt
def extract_text_from_txt(file_path):
    loader = TextLoader(file_path, encoding='utf-8')
    docs = loader.load()
    return docs


# Read File
def read_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        print("Reading PDF file:", extract_text_from_pdf(file_path))
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext in [".xlsx",".xls"]:
        return extract_text_from_excel(file_path)
    elif ext == '.csv':
        return extract_text_from_csv(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
